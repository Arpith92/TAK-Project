import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO

st.title("Itinerary Generator by TravelaajKal")

uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx"])

if uploaded_file:
    try:
        xl = pd.ExcelFile(uploaded_file)
        sheet_names = xl.sheet_names

        # Auto-select the client sheet (first one)
        client_sheet = [name for name in sheet_names if name.lower() != 'code'][0]
        client_df = xl.parse(client_sheet)
        code_df = xl.parse('Code')

        # Merge the client data with code data
        merged_df = pd.merge(client_df, code_df, on='Code', how='left')

        # Create the itinerary document
        doc = Document()
        doc.add_heading(f"Itinerary for {client_sheet}", level=1)

        for idx, row in merged_df.iterrows():
            date = row['Date']
            time = row['Time']
            code = row['Code']
            stay = row['Stay City']
            car_cost = row.get('Car Cost', 'NA')
            hotel_cost = row.get('Hotel Cost', 'NA')
            description = row.get('Description', 'No description available')

            doc.add_paragraph(
                f"📅 **{date}** at ⏰ **{time}** - Code: {code} | Stay: {stay}\n"
                f"➡ {description}\n"
                f"🚗 Car Cost: ₹{car_cost}, 🏨 Hotel Cost: ₹{hotel_cost}"
            )

        # Save to buffer
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        st.success("Itinerary generated successfully!")
        st.download_button(
            label="Download Itinerary DOCX",
            data=buf,
            file_name="Itinerary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
